# =====================================================================
# core.py - CitaLocal
# ---------------------------------------------------------------------
# MOTOR PRINCIPAL del buscador y sintetizador de literatura científica.
# No se ejecuta directamente. Es usado por app.py y literatura.py.
#
#   1) Configuración general
#   2) Utilidades internas
#   3) Comunicación con Ollama (IA local)
#   4) Conectores a bases de datos científicas (fuentes gratuitas)
#   5) Deduplicación de resultados
#   6) Orquestador de búsqueda
#   7) Generador de análisis IMRAD (síntesis con IA, dos modelos)
#   8) Exportación de archivos
#   9) Historial local de búsquedas
# =====================================================================
# Copyright (C) 2025 jbnen3
# Este programa es software libre: puedes redistribuirlo y/o modificarlo
# bajo los términos de la Licencia Pública General de GNU (GPLv3),
# publicada por la Free Software Foundation.
# Ver el archivo LICENSE para más detalles.
import os
import re
import csv
import json
import time
from datetime import datetime
import requests
import xml.etree.ElementTree as ET
from docx import Document


# =====================================================================
# 1) CONFIGURACIÓN GENERAL
# =====================================================================
NCBI_API_KEY = os.getenv("NCBI_API_KEY", "")
CORE_API_KEY = os.getenv("CORE_API_KEY", "")
OPENALEX_EMAIL = os.getenv("OPENALEX_EMAIL", "")

OLLAMA_API_URL = "http://localhost:11434"
OLLAMA_MODEL = "phi3:mini"
MAX_RESULTS = 15
PAUSA = 0.3

ANSI_RE = re.compile(r'\x1b\[[0-9;]*[a-zA-Z]')


# =====================================================================
# 2) UTILIDADES INTERNAS
# =====================================================================
def _limpiar_texto(texto):
    texto = ANSI_RE.sub('', texto)
    return texto.replace('\r', '').strip()


def _en_rango(year_str, year_from, year_to):
    try:
        y = int(str(year_str)[:4])
    except (ValueError, TypeError):
        return True
    if year_from and y < year_from:
        return False
    if year_to and y > year_to:
        return False
    return True


# =====================================================================
# 3) COMUNICACIÓN CON OLLAMA
# =====================================================================
def ollama_disponible():
    try:
        requests.get(f"{OLLAMA_API_URL}/api/tags", timeout=3)
        return True
    except Exception:
        return False


def listar_modelos_ollama():
    try:
        r = requests.get(f"{OLLAMA_API_URL}/api/tags", timeout=5)
        data = r.json()
        modelos = [m["name"] for m in data.get("models", [])]
        return modelos if modelos else [OLLAMA_MODEL]
    except Exception:
        return [OLLAMA_MODEL]
def ordenar_modelos_por_tamano():
    """Devuelve (modelo_ligero, modelo_pesado) basándose en el tamaño real."""
    try:
        r = requests.get(f"{OLLAMA_API_URL}/api/tags", timeout=5)
        modelos = r.json().get("models", [])
        if not modelos:
            return OLLAMA_MODEL, OLLAMA_MODEL
        modelos.sort(key=lambda m: m.get("size", 0))
        ligero = modelos[0]["name"]
        pesado = modelos[-1]["name"]
        return ligero, pesado
    except Exception:
        return OLLAMA_MODEL, OLLAMA_MODEL
def ordenar_modelos_por_tamano():
    """Devuelve (modelo_ligero, modelo_pesado) basándose en el tamaño real."""
    try:
        r = requests.get(f"{OLLAMA_API_URL}/api/tags", timeout=5)
        modelos = r.json().get("models", [])
        if not modelos:
            return OLLAMA_MODEL, OLLAMA_MODEL
        modelos.sort(key=lambda m: m.get("size", 0))
        ligero = modelos[0]["name"]
        pesado = modelos[-1]["name"]
        return ligero, pesado
    except Exception:
        return OLLAMA_MODEL, OLLAMA_MODEL
def ollama_generate(prompt, modelo=OLLAMA_MODEL, timeout=300, temperature=0.6):
    """
    Envía un texto (prompt) al modelo de IA local.
    temperature más baja (ej. 0.3) = respuestas más apegadas a las
    instrucciones y menos "creativas". Útil para forzar formato.
    """
    try:
        r = requests.post(
            f"{OLLAMA_API_URL}/api/generate",
            json={
                "model": modelo, "prompt": prompt, "stream": False,
                "options": {"temperature": temperature}
            },
            timeout=timeout
        )
        r.raise_for_status()
        data = r.json()
        return _limpiar_texto(data.get("response", ""))
    except requests.exceptions.ConnectionError:
        return "[Error: Ollama no está activo. Ejecuta 'ollama serve' en una terminal.]"
    except Exception as e:
        return f"[Error de Ollama: {e}]"


def resumir_con_ollama(texto, modelo=OLLAMA_MODEL):
    if not texto or len(texto.strip()) < 20:
        return "Sin resumen disponible (no hay abstract)."
    prompt = (f"Resume en español, máximo 3 líneas, lenguaje sencillo, "
              f"indicando objetivo, método y hallazgo principal:\n\n{texto}\n\nResumen:")
    return ollama_generate(prompt, modelo, timeout=90)

def clasificar_articulo(titulo, abstract, modelo=OLLAMA_MODEL):
    """
    Clasifica un artículo en un área temática usando el modelo ligero.
    Retorna una de las áreas predefinidas.
    """
    areas = [
        "Agrícola/Veterinaria/Ambiental",
        "Médica/Clínica",
        "Ciencias básicas",
        "Tecnología/Computación",
        "Educación/Social",
        "Economía",
        "Jurídica",
        "Otras áreas"
    ]
    texto = f"{titulo}. {abstract[:300]}"
    prompt = f"""Clasifica el siguiente artículo científico en UNA SOLA de estas áreas:
{chr(10).join(f"- {a}" for a in areas)}

Artículo: {texto}

Responde ÚNICAMENTE con el nombre exacto del área, sin explicación ni puntuación adicional."""

    resultado = ollama_generate(prompt, modelo, timeout=60, temperature=0.1)
    resultado = resultado.strip().strip(".-").strip()
    for area in areas:
        if area.lower() in resultado.lower():
            return area
    return "Otras áreas"


# =====================================================================
# 4) CONECTORES A BASES DE DATOS CIENTÍFICAS (solo fuentes gratuitas)
# =====================================================================
def search_pubmed(query, max_results=MAX_RESULTS, year_from=None, year_to=None, excluir=None):
    base = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"

    query_final = query
    if excluir:
        terminos_not = " NOT ".join(excluir)
        query_final = f"{query} NOT ({terminos_not})"

    params = {"db": "pubmed", "term": query_final, "retmax": max_results,
              "retmode": "json", "sort": "relevance"}

    if year_from or year_to:
        params["datetype"] = "pdat"
        params["mindate"] = str(year_from) if year_from else "1900"
        params["maxdate"] = str(year_to) if year_to else str(datetime.now().year)

    if NCBI_API_KEY:
        params["api_key"] = NCBI_API_KEY

    r = requests.get(f"{base}/esearch.fcgi", params=params, timeout=20)
    ids = r.json()["esearchresult"].get("idlist", [])
    if not ids:
        return []

    fparams = {"db": "pubmed", "id": ",".join(ids), "retmode": "xml"}
    if NCBI_API_KEY:
        fparams["api_key"] = NCBI_API_KEY
    r2 = requests.get(f"{base}/efetch.fcgi", params=fparams, timeout=20)
    return _parse_pubmed_xml(r2.text)

def _parse_pubmed_xml(xml_text):
    root = ET.fromstring(xml_text)
    results = []
    for art in root.findall(".//PubmedArticle"):
        title = art.findtext(".//ArticleTitle") or ""
        abstract = " ".join(a.text or "" for a in art.findall(".//AbstractText"))
        journal = art.findtext(".//Journal/Title") or ""
        year = art.findtext(".//PubDate/Year") or art.findtext(".//PubDate/MedlineDate") or ""
        authors = []
        for au in art.findall(".//Author"):
            last, fore = au.findtext("LastName"), au.findtext("ForeName")
            if last:
                authors.append(f"{last}, {fore or ''}".strip())
        doi = next((i.text for i in art.findall(".//ArticleId") if i.get("IdType") == "doi"), "")
        results.append({"source": "PubMed", "title": title, "abstract": abstract,
                         "journal": journal, "year": year, "authors": authors, "doi": doi})
    return results


def search_europepmc(query, max_results=MAX_RESULTS):
    url = "https://www.ebi.ac.uk/europepmc/webservices/rest/search"
    params = {"query": query, "format": "json", "pageSize": max_results}
    data = requests.get(url, params=params, timeout=20).json()
    results = []
    for it in data.get("resultList", {}).get("result", []):
        results.append({
            "source": "EuropePMC", "title": it.get("title", ""),
            "abstract": it.get("abstractText", ""), "journal": it.get("journalTitle", ""),
            "year": it.get("pubYear", ""),
            "authors": it.get("authorString", "").split(", ") if it.get("authorString") else [],
            "doi": it.get("doi", "")
        })
    return results


def search_crossref(query, max_results=MAX_RESULTS, year_from=None, year_to=None):
    url = "https://api.crossref.org/works"
    params = {"query": query, "rows": max_results}
    if year_from or year_to:
        f = []
        if year_from:
            f.append(f"from-pub-date:{year_from}-01-01")
        if year_to:
            f.append(f"until-pub-date:{year_to}-12-31")
        params["filter"] = ",".join(f)
    data = requests.get(url, params=params, timeout=20).json()
    results = []
    for it in data.get("message", {}).get("items", []):
        title = it.get("title", [""])[0] if it.get("title") else ""
        abstract = re.sub("<[^<]+?>", "", it.get("abstract", "")) if it.get("abstract") else ""
        authors = [f"{a.get('family','')}, {a.get('given','')}" for a in it.get("author", [])]
        year = ""
        for k in ("published-print", "published-online"):
            if k in it:
                year = str(it[k]["date-parts"][0][0]); break
        results.append({
            "source": "Crossref", "title": title, "abstract": abstract,
            "journal": (it.get("container-title") or [""])[0], "year": year,
            "authors": authors, "doi": it.get("DOI", "")
        })
    return results


def search_semanticscholar(query, max_results=MAX_RESULTS):
    url = "https://api.semanticscholar.org/graph/v1/paper/search"
    params = {"query": query, "limit": max_results,
              "fields": "title,abstract,year,venue,authors,externalIds"}
    data = requests.get(url, params=params, timeout=20).json()
    results = []
    for it in data.get("data", []):
        results.append({
            "source": "SemanticScholar", "title": it.get("title", ""),
            "abstract": it.get("abstract") or "", "journal": it.get("venue", ""),
            "year": str(it.get("year", "")),
            "authors": [a["name"] for a in it.get("authors", [])],
            "doi": it.get("externalIds", {}).get("DOI", "")
        })
    return results


def search_doaj(query, max_results=MAX_RESULTS):
    url = f"https://doaj.org/api/v3/search/articles/{requests.utils.quote(query)}"
    params = {"pageSize": max_results}
    try:
        data = requests.get(url, params=params, timeout=15).json()
    except Exception:
        return []
    results = []
    for it in data.get("results", []):
        bib = it.get("bibjson", {})
        authors = [a.get("name", "") for a in bib.get("author", [])]
        doi = next((i.get("id") for i in bib.get("identifier", [])
                    if i.get("type") == "doi"), "")
        results.append({"source": "DOAJ/MDPI", "title": bib.get("title", ""),
                         "abstract": bib.get("abstract", ""),
                         "journal": bib.get("journal", {}).get("title", ""),
                         "year": str(bib.get("year", "")), "authors": authors, "doi": doi})
    return results


def _reconstruir_abstract(inverted_index):
    if not inverted_index:
        return ""
    posiciones = {}
    max_pos = 0
    for palabra, pos_list in inverted_index.items():
        for p in pos_list:
            posiciones[p] = palabra
            if p > max_pos:
                max_pos = p
    return " ".join(posiciones.get(i, "") for i in range(max_pos + 1))


def search_openalex(query, max_results=MAX_RESULTS, year_from=None, year_to=None):
    url = "https://api.openalex.org/works"
    params = {"search": query, "per-page": max_results}
    if OPENALEX_EMAIL:
        params["mailto"] = OPENALEX_EMAIL
    filtros = []
    if year_from:
        filtros.append(f"from_publication_date:{year_from}-01-01")
    if year_to:
        filtros.append(f"to_publication_date:{year_to}-12-31")
    if filtros:
        params["filter"] = ",".join(filtros)
    try:
        data = requests.get(url, params=params, timeout=20).json()
    except Exception:
        return []
    results = []
    for it in data.get("results", []):
        abstract = _reconstruir_abstract(it.get("abstract_inverted_index"))
        authors = [a["author"]["display_name"] for a in it.get("authorships", [])
                   if a.get("author")]
        primary = it.get("primary_location") or {}
        source = primary.get("source") or {}
        venue = source.get("display_name", "") if source else ""
        doi = (it.get("doi") or "").replace("https://doi.org/", "")
        results.append({
            "source": "OpenAlex", "title": it.get("title") or "",
            "abstract": abstract, "journal": venue,
            "year": str(it.get("publication_year", "")),
            "authors": authors, "doi": doi
        })
    return results


def search_arxiv(query, max_results=MAX_RESULTS):
    url = "http://export.arxiv.org/api/query"
    params = {"search_query": f"all:{query}", "max_results": max_results}
    try:
        r = requests.get(url, params=params, timeout=20)
    except Exception:
        return []
    ns = {"atom": "http://www.w3.org/2005/Atom"}
    try:
        root = ET.fromstring(r.text)
    except Exception:
        return []
    results = []
    for entry in root.findall("atom:entry", ns):
        title = (entry.findtext("atom:title", default="", namespaces=ns) or "").strip().replace("\n", " ")
        summary = (entry.findtext("atom:summary", default="", namespaces=ns) or "").strip().replace("\n", " ")
        published = entry.findtext("atom:published", default="", namespaces=ns) or ""
        authors = [a.findtext("atom:name", default="", namespaces=ns)
                   for a in entry.findall("atom:author", ns)]
        results.append({"source": "arXiv", "title": title, "abstract": summary,
                         "journal": "arXiv (preprint)", "year": published[:4],
                         "authors": authors, "doi": ""})
    return results


def search_core(query, max_results=MAX_RESULTS):
    if not CORE_API_KEY:
        return []
    url = "https://api.core.ac.uk/v3/search/works"
    headers = {"Authorization": f"Bearer {CORE_API_KEY}"}
    params = {"q": query, "limit": max_results}
    try:
        data = requests.get(url, headers=headers, params=params, timeout=20).json()
    except Exception:
        return []
    results = []
    for it in data.get("results", []):
        authors = []
        for a in it.get("authors", []) or []:
            if isinstance(a, dict):
                authors.append(a.get("name", ""))
            elif isinstance(a, str):
                authors.append(a)
        results.append({
            "source": "CORE", "title": it.get("title", "") or "",
            "abstract": it.get("abstract", "") or "",
            "journal": it.get("publisher", "") or "",
            "year": str(it.get("yearPublished") or ""),
            "authors": authors, "doi": it.get("doi") or ""
        })
    return results

# =====================================================================
# AGRIS (FAO) — Base de datos agrícola internacional
# =====================================================================
def search_agris(query, max_results=MAX_RESULTS):
    url = "https://agris.fao.org/agris-search/search.do"
    params = {
        "query": query,
        "startRow": 0,
        "format": "json",
        "rows": max_results
    }
    try:
        data = requests.get(url, params=params, timeout=15).json()
    except Exception:
        return []
    results = []
    for it in data.get("response", {}).get("docs", []):
        authors = it.get("creatorPersonal", []) or []
        results.append({
            "source": "AGRIS/FAO",
            "title": it.get("title", "") or "",
            "abstract": it.get("abstract", "") or "",
            "journal": it.get("hostTitle", "") or "",
            "year": str(it.get("publicationDate", "") or "")[:4],
            "authors": authors if isinstance(authors, list) else [authors],
            "doi": it.get("identifier", "") or ""
        })
    return results


# =====================================================================
# AGRICOLA (USDA) — Base de datos agrícola de EE.UU.
# =====================================================================
def search_agricola(query, max_results=MAX_RESULTS):
    url = "https://api.nal.usda.gov/agricola/search"
    params = {
        "query": query,
        "max": max_results,
        "format": "json"
    }
    try:
        data = requests.get(url, params=params, timeout=15).json()
    except Exception:
        return []
    results = []
    for it in data.get("result", []):
        results.append({
            "source": "AGRICOLA/USDA",
            "title": it.get("title", "") or "",
            "abstract": it.get("abstract", "") or "",
            "journal": it.get("journal", "") or "",
            "year": str(it.get("year", "") or "")[:4],
            "authors": it.get("author", []) or [],
            "doi": it.get("doi", "") or ""
        })
    return results


# =====================================================================
# SSRN — Ciencias sociales y economía
# =====================================================================
def search_ssrn(query, max_results=MAX_RESULTS):
    url = "https://api.ssrn.com/content/search"
    params = {
        "query": query,
        "limit": max_results
    }
    try:
        data = requests.get(url, params=params, timeout=15).json()
    except Exception:
        return []
    results = []
    for it in data.get("papers", []):
        authors = [a.get("name", "") for a in it.get("authors", [])]
        results.append({
            "source": "SSRN",
            "title": it.get("title", "") or "",
            "abstract": it.get("abstract", "") or "",
            "journal": "SSRN",
            "year": str(it.get("date", "") or "")[:4],
            "authors": authors,
            "doi": it.get("doi", "") or ""
        })
    return results


# =====================================================================
# ERIC — Educación
# =====================================================================
def search_eric(query, max_results=MAX_RESULTS):
    url = "https://api.ies.ed.gov/eric/"
    params = {
        "search": query,
        "format": "json",
        "rows": max_results,
        "fields": "title,author,publicationdateyear,description,issn,doi"
    }
    try:
        data = requests.get(url, params=params, timeout=15).json()
    except Exception:
        return []
    results = []
    for it in data.get("response", {}).get("docs", []):
        authors = it.get("author", []) or []
        results.append({
            "source": "ERIC/Educación",
            "title": it.get("title", "") or "",
            "abstract": it.get("description", "") or "",
            "journal": it.get("issn", "") or "",
            "year": str(it.get("publicationdateyear", "") or ""),
            "authors": authors if isinstance(authors, list) else [authors],
            "doi": it.get("doi", "") or ""
        })
    return results
# =====================================================================
# UNPAYWALL — Verificador de acceso abierto por DOI
# =====================================================================
UNPAYWALL_EMAIL = os.getenv("UNPAYWALL_EMAIL", "claudio.bernal@uabc.edu.mx")

def verificar_unpaywall(doi):
    """
    Dado un DOI, consulta Unpaywall para saber si existe
    una versión gratuita y legal del artículo completo.
    Retorna la URL de acceso abierto o None si no existe.
    """
    if not doi:
        return None
    try:
        url = f"https://api.unpaywall.org/v2/{doi}"
        params = {"email": UNPAYWALL_EMAIL}
        r = requests.get(url, params=params, timeout=10)
        if r.status_code != 200:
            return None
        data = r.json()
        if data.get("is_oa"):
            best = data.get("best_oa_location") or {}
            return best.get("url_for_pdf") or best.get("url")
        return None
    except Exception:
        return None

TODAS_LAS_FUENTES = {
    "PubMed": search_pubmed,
    "EuropePMC": search_europepmc,
    "Crossref": search_crossref,
    "SemanticScholar": search_semanticscholar,
    "DOAJ/MDPI": search_doaj,
    "OpenAlex": search_openalex,
    "arXiv": search_arxiv,
    "CORE": search_core,
    "AGRIS/FAO": search_agris,
    "AGRICOLA/USDA": search_agricola,
    "SSRN": search_ssrn,
    "ERIC/Educación": search_eric,
}

_FUENTES_CON_FILTRO_ANIO = ("PubMed", "Crossref", "OpenAlex")


# =====================================================================
# 5) DEDUPLICACIÓN
# =====================================================================
def deduplicar(resultados):
    vistos, unicos = set(), []
    for r in resultados:
        clave = r["doi"].lower() if r["doi"] else r["title"].lower().strip()[:60]
        if clave and clave not in vistos:
            vistos.add(clave)
            unicos.append(r)
    return unicos


# =====================================================================
# 6) ORQUESTADOR DE BÚSQUEDA
# =====================================================================
def buscar_literatura(query, max_results=MAX_RESULTS, fuentes_activas=None,
                       year_from=None, year_to=None,
                       generar_resumen=True, modelo=OLLAMA_MODEL,
                       excluir=None, progreso_callback=None):
    if fuentes_activas is None:
        fuentes_activas = list(TODAS_LAS_FUENTES.keys())

    def log(msg, frac=None):
        if progreso_callback:
            progreso_callback(msg, frac)

    todos_raw = []
    for nombre in fuentes_activas:
        fn = TODAS_LAS_FUENTES.get(nombre)
        if not fn:
            continue
        log(f"Buscando en {nombre}...")
        try:
            if nombre in _FUENTES_CON_FILTRO_ANIO:
                res = fn(query, max_results, year_from, year_to)
            else:
                res = fn(query, max_results)
            todos_raw += res
        except Exception as e:
            log(f"{nombre} falló: {e}")
        time.sleep(PAUSA)

    todos = deduplicar(todos_raw)
    if excluir:
        todos = [
            r for r in todos
            if not any(
                term.lower() in (r.get("title", "") + " " + r.get("abstract", "")).lower()
                for term in excluir
            )
        ]
    todos = [r for r in todos if _en_rango(r["year"], year_from, year_to)]

    for i, r in enumerate(todos):
        log(f"Clasificando y procesando [{i+1}/{len(todos)}]: {r['title'][:50]}")
        r["area"] = clasificar_articulo(r.get("title", ""), r.get("abstract", ""), modelo)
        r["url_oa"] = verificar_unpaywall(r.get("doi", ""))
    log("Listo.")

    return todos


# =====================================================================
# 7) ANÁLISIS IMRAD (dos etapas, dos modelos posibles)
# ---------------------------------------------------------------------
# Etapa A (map): extrae puntos clave por lotes -> puede usar un modelo
#                rápido y pequeño, aquí sí está bien que use viñetas
#                porque es solo material de trabajo intermedio.
# Etapa B (reduce): redacta el texto final -> se recomienda un modelo
#                más grande/capaz, y se fuerza estilo narrativo con
#                reglas estrictas para evitar que copie las viñetas.
# =====================================================================
PLANTILLA_NARRATIVA = """Eres un investigador científico senior y redactor académico en español con 20 años de experiencia. Tu tarea es redactar una síntesis narrativa en formato IMRAD sobre "{tema}".

REGLAS OBLIGATORIAS — incumplirlas invalida tu respuesta:
1. ESTILO: Escribe ÚNICAMENTE en prosa académica fluida, en párrafos completos y conectados.
2. PROHIBIDO: viñetas, guiones, asteriscos, numeración, listas de cualquier tipo.
3. CITAS: cada afirmación importante debe ir acompañada de su cita en formato (Apellido, año). Ejemplo: "diversos estudios han demostrado que el injerto mejora la resistencia (García, 2020; López, 2022)".
4. AGRUPACIÓN: cuando varios autores dicen lo mismo, agrúpalos en una sola cita: (Autor1, año; Autor2, año).
5. TRANSICIONES: conecta ideas con "asimismo", "por otro lado", "en este sentido", "cabe destacar que", "en contraste", "estos hallazgos sugieren que".
6. SÍNTESIS: no resumas artículo por artículo — integra tendencias, consensos y contradicciones.
7. PRECISIÓN: usa solo datos del material de referencia. No inventes autores ni cifras.
8. EXTENSIÓN: mínimo 3 párrafos por sección, 600-800 palabras en total.

Usa exactamente estos encabezados y ningún otro:
## Introducción
## Métodos
## Resultados
## Discusión

Material de referencia con autores y años (usa estos datos para las citas):
{combinado}

Escribe ahora la síntesis completa en párrafos narrativos con citas APA:"""
PLANTILLA_ESTRUCTURADA = """Eres un redactor académico experto en español. Organiza una síntesis estructurada sobre "{tema}" usando los siguientes encabezados exactos:

## Introducción
## Métodos
## Resultados
## Discusión

REGLAS:
1. Dentro de cada sección escribe primero 1-2 párrafos narrativos de contexto.
2. Luego puedes usar viñetas SOLO para listar hallazgos específicos con datos concretos.
3. Cada viñeta debe ser una oración completa con sujeto, verbo y dato.
4. Conecta las secciones con frases de transición entre ellas.
5. PROHIBIDO: repetir la misma idea en secciones diferentes.

Material de referencia:
{combinado}

Escribe la síntesis estructurada:"""


def generar_analisis_imrad(resultados, tema,
                            modelo_lotes=OLLAMA_MODEL, modelo_final=None,
                            progreso_callback=None, tamano_lote=6,
                            estilo="narrativo", prompt_personalizado=None):
    """
    modelo_lotes  -> modelo usado para extraer puntos clave por lotes (rápido)
    modelo_final  -> modelo usado para redactar el texto final (calidad).
                     Si es None, se usa el mismo que modelo_lotes.
    estilo        -> "narrativo" (párrafos) o "estructurado" (viñetas)
    prompt_personalizado -> si el usuario escribe su propio prompt, debe
                     incluir los marcadores {tema} y {combinado}
    """
    if not resultados:
        return "No hay artículos para analizar."
    if modelo_final is None:
        modelo_final = modelo_lotes

    def log(msg, frac=None):
        if progreso_callback:
            progreso_callback(msg, frac)

    # --- Etapa A: análisis por lotes ---
    lotes = [resultados[i:i + tamano_lote] for i in range(0, len(resultados), tamano_lote)]
    total_pasos = len(lotes) + 1
    sintesis_parciales = []

    for idx, lote in enumerate(lotes):
        texto_lote = "\n\n".join(
            f"- {r['title']} ({', '.join(r['authors'][:2]) if r['authors'] else 'Autor desconocido'}, {r['year']}): {r.get('resumen') or r['abstract'][:400]}"
            for r in lote
        )
        prompt_lote = f"""Eres un asistente de investigación. A partir de los siguientes artículos científicos sobre "{tema}", extrae en español y en viñetas breves:
1) Objetivos comunes de las investigaciones
2) Métodos o diseños de estudio empleados
3) Principales resultados o hallazgos
4) Puntos de discusión, coincidencias o controversias

No inventes información que no esté en los textos.

Artículos:
{texto_lote}
"""
        log(f"Analizando lote {idx+1}/{len(lotes)} (modelo: {modelo_lotes})...", idx / total_pasos)
        resultado = ollama_generate(prompt_lote, modelo_lotes, timeout=180, temperature=0.5)
        sintesis_parciales.append(resultado)

    combinado = "\n\n---\n\n".join(sintesis_parciales)

    # --- Etapa B: redacción final ---
    if prompt_personalizado and prompt_personalizado.strip():
        prompt_final = prompt_personalizado.format(tema=tema, combinado=combinado)
    elif estilo == "estructurado":
        prompt_final = PLANTILLA_ESTRUCTURADA.format(tema=tema, combinado=combinado)
    else:
        prompt_final = PLANTILLA_NARRATIVA.format(tema=tema, combinado=combinado)

    log(f"Generando síntesis final (modelo: {modelo_final})...", len(lotes) / total_pasos)
    # Temperatura baja para que respete mejor las reglas de formato.
    texto_final = ollama_generate(prompt_final, modelo_final, timeout=300, temperature=0.20)
    log("Análisis completado.", 1.0)
    return texto_final


# =====================================================================
# 8) EXPORTACIÓN DE ARCHIVOS
# =====================================================================
def exportar_ris(resultados, filename="literatura.ris"):
    with open(filename, "w", encoding="utf-8") as f:
        for r in resultados:
            f.write("TY  - JOUR\n")
            f.write(f"TI  - {r['title']}\n")
            for a in r["authors"]:
                f.write(f"AU  - {a}\n")
            f.write(f"PY  - {r['year']}\nJO  - {r['journal']}\n")
            if r["doi"]:
                f.write(f"DO  - {r['doi']}\n")
            f.write(f"AB  - {r['abstract']}\nER  - \n\n")


def exportar_bib(resultados, filename="literatura.bib"):
    with open(filename, "w", encoding="utf-8") as f:
        for r in resultados:
            autor = re.sub(r"[^a-zA-Z]", "", r["authors"][0].split(",")[0]) if r["authors"] else "anon"
            key = f"{autor}{r['year']}"
            autores = " and ".join(r["authors"]) if r["authors"] else "Unknown"
            f.write(f"@article{{{key},\n  title = {{{r['title']}}},\n"
                     f"  author = {{{autores}}},\n  year = {{{r['year']}}},\n"
                     f"  journal = {{{r['journal']}}},\n")
            if r["doi"]:
                f.write(f"  doi = {{{r['doi']}}},\n")
            f.write(f"  abstract = {{{r['abstract']}}}\n}}\n\n")


def exportar_csv(resultados, filename="literatura.csv"):
   with open(filename, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["Fuente", "Autores", "Año", "Título", "Revista", "DOI", "Acceso Abierto", "Resumen"])
        for r in resultados:
            w.writerow([r["source"], "; ".join(r["authors"]), r["year"],
                        r["title"], r["journal"], r["doi"],
                        r.get("url_oa", "") or "No disponible",
                        r.get("resumen", "")])

def exportar_markdown(resultados, filename="resumenes.md", query=""):
    with open(filename, "w", encoding="utf-8") as f:
        f.write(f"# Revisión de literatura: {query}\n\n")
        f.write(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")
        for r in resultados:
            f.write(f"## {r['title']}\n")
            f.write(f"**Fuente:** {r['source']} | **Autores:** {'; '.join(r['authors'])}\n\n")
            f.write(f"**Año:** {r['year']} | **Revista:** {r['journal']} | **DOI:** {r['doi']} | **Acceso Abierto:** {r.get('url_oa', '') or 'No disponible'}\n\n")
            f.write(f"**Resumen:** {r.get('resumen','N/A')}\n\n---\n\n")


def exportar_analisis_md(texto, filename="analisis_imrad.md", tema=""):
    with open(filename, "w", encoding="utf-8") as f:
        f.write(f"# Síntesis de literatura (IMRAD): {tema}\n\n")
        f.write(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")
        f.write(texto)


def exportar_analisis_txt(texto, filename="analisis_imrad.txt", tema=""):
    with open(filename, "w", encoding="utf-8") as f:
        f.write(f"SÍNTESIS DE LITERATURA (IMRAD): {tema}\n")
        f.write(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")
        f.write(texto.replace("## ", "").replace("# ", ""))


def exportar_analisis_docx(texto, filename="analisis_imrad.docx", tema=""):
    doc = Document()
    doc.add_heading(f"Síntesis de literatura: {tema}", level=1)
    doc.add_paragraph(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    for linea in texto.split("\n"):
        linea = linea.strip()
        if not linea:
            continue
        if linea.startswith("## "):
            doc.add_heading(linea.replace("## ", ""), level=2)
        elif linea.startswith("# "):
            doc.add_heading(linea.replace("# ", ""), level=1)
        elif linea.startswith("- ") or linea.startswith("* "):
            doc.add_paragraph(linea[2:], style="List Bullet")
        else:
            doc.add_paragraph(linea)
    doc.save(filename)


# =====================================================================
# 9) HISTORIAL LOCAL DE BÚSQUEDAS
# =====================================================================
HISTORIAL_DIR = "historial"


def guardar_busqueda(tema, resultados, year_from=None, year_to=None, carpeta=HISTORIAL_DIR):
    os.makedirs(carpeta, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    slug = re.sub(r'[^a-zA-Z0-9]+', '_', tema.strip().lower())[:40].strip('_') or "busqueda"
    filename = os.path.join(carpeta, f"{timestamp}_{slug}.json")
    payload = {
        "tema": tema, "fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "year_from": year_from, "year_to": year_to,
        "total": len(resultados), "resultados": resultados,
    }
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    return filename


def listar_busquedas_recientes(carpeta=HISTORIAL_DIR, limite=15):
    if not os.path.isdir(carpeta):
        return []
    archivos = sorted(
        [os.path.join(carpeta, f) for f in os.listdir(carpeta) if f.endswith(".json")],
        key=os.path.getmtime, reverse=True
    )[:limite]
    items = []
    for path in archivos:
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            items.append({"path": path, "tema": data.get("tema", "(sin tema)"),
                          "fecha": data.get("fecha", ""), "total": data.get("total", 0)})
        except Exception:
            continue
    return items


def cargar_busqueda(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)